import io
import csv
import logging
from pypdf import PdfReader
from fastapi import UploadFile

logger = logging.getLogger("agent.extractor")

# Guard the python-docx import
try:
    from docx import Document
except ImportError:
    logger.error("python-docx is not installed or failed to import.")
    Document = None

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text programmatically using pypdf safely."""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        extracted_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                extracted_text.append(text)
        return "\n".join(extracted_text).strip()
    except Exception as e:
        logger.error(f"Failed to parse PDF: {str(e)}")
        return f"[Error parsing PDF: {str(e)}]"

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text safely from Word documents, including paragraphs and tables."""
    if Document is None:
        return "[Error: Word document parser (python-docx) is not installed on the server.]"
    
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        extracted_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_text.append(" | ".join(row_text))
                    
        return "\n".join(extracted_text).strip()
    except Exception as e:
        logger.error(f"Failed to parse Word document: {str(e)}")
        return f"[Error parsing Word Document: {str(e)}]"

def extract_text_from_csv(file_bytes: bytes) -> str:
    """Formats CSV rows into a readable structured text dataset safely."""
    try:
        csv_file = io.StringIO(file_bytes.decode('utf-8', errors='ignore'))
        reader = csv.reader(csv_file)
        rows_text = []
        for row in reader:
            rows_text.append(" | ".join(row))
        return "\n".join(rows_text).strip()
    except Exception as e:
        logger.error(f"Failed to parse CSV: {str(e)}")
        return f"[Error parsing CSV: {str(e)}]"

async def extract_text_from_image(file_bytes: bytes, mime_type: str) -> str:
    """Performs native multimodal OCR on image bytes using Gemini."""
    try:
        from google import genai
        from google.genai import types
        from app.config import settings
        from app.utils.retry import execute_with_retry
        
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        
        prompt = (
            "You are an expert OCR engine. Analyze this image and perform the following:\n"
            "1. Extract all visible text exactly as it appears. Keep layout and spacing clean.\n"
            "2. Estimate your overall OCR confidence score as a percentage (e.g., 95%).\n\n"
            "Respond exactly in this format:\n"
            "--- Cleaned Transcript ---\n"
            "[Extracted text here]\n\n"
            "--- OCR Confidence ---\n"
            "[Confidence percentage, e.g., 98%]"
        )
        
        response = await execute_with_retry(
            client.aio.models.generate_content,
            model=settings.GEMINI_MODEL,
            contents=[
                types.Part.from_bytes(data=file_bytes, mime_type=mime_type),
                prompt
            ]
        )
        return response.text.strip()
    except Exception as e:
        logger.error(f"Failed to perform OCR on image: {str(e)}")
        return f"[Error performing OCR on image: {str(e)}]"

async def extract_text_from_audio(file_bytes: bytes, mime_type: str) -> str:
    """
    Transcribes audio and extracts duration using a robust Hybrid Engine:
    1. Attempts local metadata header extraction (TinyTag via tempfile).
    2. Falls back to native Gemini AI duration parsing if local extraction fails,
       utilizing strict break boundary guards to prevent transcript overwrites.
    """
    import os
    import tempfile
    from tinytag import TinyTag
    
    # Method A: Try local metadata header extraction first (extremely fast & accurate)
    local_duration_str = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".tmp") as temp_file:
            temp_file.write(file_bytes)
            temp_file_path = temp_file.name
        
        tag = TinyTag.get(temp_file_path)
        duration_sec = tag.duration or 0.0
        os.remove(temp_file_path)
        
        if duration_sec > 0:
            minutes = int(duration_sec // 60)
            seconds = int(duration_sec % 60)
            local_duration_str = f"{minutes} min {seconds} sec"
    except Exception as e:
        logger.warning(f"Local TinyTag extraction bypassed (using AI fallback): {str(e)}")

    # Method B: Transcribe and perform AI duration extraction with strict loop boundaries
    try:
        from google import genai
        from google.genai import types
        from app.config import settings
        from app.utils.retry import execute_with_retry
        
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        
        prompt = (
            "You are an expert audio analyzer and transcriber. Listen to this audio and perform the following:\n"
            "1. Transcribe all spoken words verbatim, cleaning up filler words (um, ah, like).\n"
            "2. Determine the exact duration of this audio file (e.g., '1 min 45 sec' or '5 min 0 sec').\n\n"
            "You MUST respond strictly using this template format:\n"
            "DURATION: [Insert duration here]\n"
            "TRANSCRIPT:\n"
            "[Insert transcription here]"
        )
        
        response = await execute_with_retry(
            client.aio.models.generate_content,
            model=settings.GEMINI_MODEL,
            contents=[
                types.Part.from_bytes(data=file_bytes, mime_type=mime_type),
                prompt
            ]
        )
        
        response_text = response.text.strip()
        
        # Parse AI transcription
        transcript_lines = []
        is_reading_transcript = False
        ai_duration_str = ""
        
        for line in response_text.split("\n"):
            clean_line = line.strip().replace("*", "")
            # Check strictly for the start of the line and break immediately once found
            if clean_line.lower().startswith("duration:"):
                ai_duration_str = clean_line.split(":", 1)[1].strip()
            elif clean_line.lower().startswith("transcript:"):
                is_reading_transcript = True
            elif is_reading_transcript:
                transcript_lines.append(line)
        
        transcript_text = "\n".join(transcript_lines).strip()
        if not transcript_text:
            transcript_text = response_text
            
        # Select the best available duration value (Prioritize local, fallback to AI)
        final_duration = local_duration_str if local_duration_str else (ai_duration_str if ai_duration_str else "Unknown duration")
        
        return (
            f"--- Audio Transcription ---\n{transcript_text}\n\n"
            f"--- Audio Duration ---\n{final_duration}"
        )
    except Exception as e:
        logger.error(f"Failed to transcribe audio: {str(e)}")
        return f"[Error transcribing audio: {str(e)}]"

async def extract_text_from_file(file: UploadFile) -> str:
    """Dispatches the correct extractor safely based on file extension, including Audio."""
    try:
        content = await file.read()
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            return extract_text_from_pdf(content)
        elif filename.endswith(".docx"):
            return extract_text_from_docx(content)
        elif filename.endswith(".csv"):
            return extract_text_from_csv(content)
        elif filename.endswith(".txt"):
            # Natively decode plain text files safely clearing non-unicode anomalies
            return content.decode("utf-8", errors="ignore").strip()
        elif filename.endswith(".png"):
            return await extract_text_from_image(content, "image/png")
        elif filename.endswith((".jpg", ".jpeg")):
            return await extract_text_from_image(content, "image/jpeg")
        elif filename.endswith(".mp3"):
            return await extract_text_from_audio(content, "audio/mp3")
        elif filename.endswith(".wav"):
            return await extract_text_from_audio(content, "audio/wav")
        elif filename.endswith(".m4a"):
            return await extract_text_from_audio(content, "audio/m4a")
        return ""
    except Exception as e:
        logger.error(f"Error in extract_text_from_file for {file.filename}: {str(e)}")
        return f"[Error extracting text from {file.filename}: {str(e)}]"